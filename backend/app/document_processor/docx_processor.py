import docx
import logging

logger = logging.getLogger(__name__)


class DOCXProcessor:
    """Extracts text from Word documents (.docx) into structured sections."""
    
    @staticmethod
    def extract(file_path: str) -> list[dict]:
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            logger.warning(f"python-docx failed on {file_path} ({e}). Attempting legacy .doc text stream extraction...")
            sections = DOCXProcessor._extract_legacy_doc_strings(file_path)
            if sections:
                return sections
            raise ValueError(f"Failed to read Word document (.docx/.doc): {str(e)}")
        
        sections = []
        current_heading = "Document Content"
        current_block = []
        current_length = 0
        section_idx = 1

        def flush_block(heading_name):
            nonlocal current_block, current_length, section_idx
            if current_block:
                body = "\n".join(current_block).strip()
                if body:
                    sections.append({
                        'page_number': section_idx,
                        'section': heading_name,
                        'text': body,
                        'file_type': 'docx'
                    })
                    section_idx += 1
                current_block = []
                current_length = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check for Heading style or bold standalone title
            is_heading = para.style and para.style.name.startswith('Heading')
            
            if is_heading:
                flush_block(current_heading)
                current_heading = text
            else:
                current_block.append(text)
                current_length += len(text)
                # Flush block if length exceeds 1000 chars to ensure good vector chunking
                if current_length >= 1000:
                    flush_block(current_heading)

        # Include tables text
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    table_rows.append(row_text)
            if table_rows:
                current_block.append("\n".join(table_rows))
                current_length += sum(len(r) for r in table_rows)
                if current_length >= 1000:
                    flush_block(f"{current_heading} (Table Data)")

        flush_block(current_heading)

        # Fallback: Check for embedded images (e.g. Canva/image resumes exported as .docx)
        if not sections or sum(len(s['text']) for s in sections) < 50:
            logger.info(f"DOCX {file_path} text is empty/sparse. Attempting embedded image OCR extraction...")
            try:
                import zipfile
                import os
                from app.services.ocr_service import OCRService

                with zipfile.ZipFile(file_path, 'r') as z:
                    media_files = [f for f in z.namelist() if f.startswith('word/media/')]
                    for img_file in media_files:
                        img_bytes = z.read(img_file)
                        ext = img_file.split('.')[-1].lower()
                        mime = 'image/jpeg' if ext in ['jpg', 'jpeg'] else 'image/png'
                        ocr_text = OCRService.extract_text_from_image_bytes(img_bytes, mime)
                        if ocr_text.strip():
                            sections.append({
                                'page_number': len(sections) + 1,
                                'section': f'Embedded Image Content ({os.path.basename(img_file)})',
                                'text': ocr_text.strip(),
                                'file_type': 'docx'
                            })
            except Exception as err:
                logger.warning(f"Failed embedded image OCR in DOCX: {err}")

        if not sections:
            # Fallback if docx has text in shapes/textboxes or unrecognized elements
            full_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            if full_text.strip():
                sections.append({
                    'page_number': 1,
                    'section': 'Full Document Text',
                    'text': full_text.strip(),
                    'file_type': 'docx'
                })
            else:
                raise ValueError("No extractable text found in Word document. Please ensure file contains clear text or images.")
            
        logger.info(f"Extracted {len(sections)} sections from Word document {file_path}")
        return sections

    @staticmethod
    def _extract_legacy_doc_strings(file_path: str) -> list[dict]:
        import re
        import os
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            raw_text = content.decode('utf-8', errors='ignore')
            matches = re.findall(r'[\x20-\x7E\t\n\r]{6,}', raw_text)
            clean_lines = []
            for m in matches:
                m_strip = m.strip()
                if len(m_strip) >= 8 and not m_strip.startswith(('Root Entry', 'WordDocument', 'Table', 'SummaryInformation', 'CompObj', 'ObjectPool')):
                    clean_lines.append(m_strip)
            
            full_text = "\n".join(clean_lines).strip()
            if len(full_text) > 20:
                logger.info(f"Successfully extracted {len(full_text)} chars from legacy Word .doc: {os.path.basename(file_path)}")
                return [{
                    'page_number': 1,
                    'section': 'Legacy Word Document (.doc) Content',
                    'text': full_text,
                    'file_type': 'doc'
                }]
        except Exception as err:
            logger.warning(f"Legacy .doc extraction error on {file_path}: {err}")
        return []
